import openai
import streamlit as st
import os
from streamlit_chat import message as msg
import docx
import io


openai.api_key = st.secrets["OPENAI_API_KEY"]

st.title("AI Assistant")
st.write("***") # line to seperate title

if 'hst_conversa' not in st.session_state: # saves data
    st.session_state.hst_conversa = []


userQuestion = st.text_area("Type the question here...")
btn = st.button("Send Question")
if btn:
    st.session_state.hst_conversa.append({"role": "user", "content": userQuestion})
    return_openai = openai.ChatCompletion.create(
        model = "gpt-3.5-turbo",
        messages = st.session_state.hst_conversa,
        max_tokens = 3000,
        n=1 # how many answers i want back
    )
    st.session_state.hst_conversa.append(
        {"role": "assistant",
         "content":return_openai['choices'][0]['message']['content']})

if len(st.session_state.hst_conversa) > 0: #if there is already history of conversation
    for i in range(len(st.session_state.hst_conversa)):
        if i % 2 == 0: # If the condition is even...
            msg("You: " + st.session_state.hst_conversa[i]['content'], is_user=True)
        else:
            msg("AI Answer: " + st.session_state.hst_conversa[i]['content'])

if len(st.session_state.hst_conversa) > 0:
     btn_save = st.button("Save Content")
     if btn_save: #if button clicked
         sWork = io.BytesIO()
         document = docx.Document()
         document.add_heading('First Heading', level=1)

         for i in range(len(st.session_state.hst_conversa)):
             if i % 2 == 0:
                document.add_heading("Question", level=2) # create heading for word doc
                document.add_paragraph(st.session_state.hst_conversa[i]['content'])
             else:
                document.add_heading("Answer", level=2)
                document.add_paragraph(st.session_state.hst_conversa[i]['content'])


         document.save(sWork) # button for saving chat text as word doc
         st.download_button(label="Click here to save content",
                           data=sWork,
                           file_name="",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        
